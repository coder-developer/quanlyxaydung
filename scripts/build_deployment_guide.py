from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from pathlib import Path

OUT = Path('release/Huong-dan-cai-dat-CONSTRUCT-OS.docx')
OUT.parent.mkdir(parents=True, exist_ok=True)
doc = Document()
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(.492)

styles = doc.styles
normal = styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(11)
normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
for name,size,before,after in [('Heading 1',16,18,10),('Heading 2',13,14,7),('Heading 3',12,10,5)]:
    s=styles[name]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor(46,116,181); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)

header=sec.header.paragraphs[0]; header.text='CONSTRUCT-OS ERP  |  Hướng dẫn triển khai'; header.runs[0].font.size=Pt(9); header.runs[0].font.color.rgb=RGBColor(100,110,120)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT
r=footer.add_run('Tài liệu vận hành nội bộ  |  '); r.font.size=Pt(9); r.font.color.rgb=RGBColor(100,110,120)
fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); footer._p.append(fld)

def title(text,size=28,color=(15,23,42),after=8):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(after)
    r=p.add_run(text); r.bold=True; r.font.name='Calibri'; r.font.size=Pt(size); r.font.color.rgb=RGBColor(*color); return p
def para(text,bold=False,italic=False,align=None):
    p=doc.add_paragraph(); p.alignment=align or WD_ALIGN_PARAGRAPH.LEFT; r=p.add_run(text); r.bold=bold; r.italic=italic; return p
def bullet(text): doc.add_paragraph(text, style='List Bullet')
def step(text): doc.add_paragraph(text, style='List Number')
def code(text):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.25); p.paragraph_format.right_indent=Inches(.25)
    r=p.add_run(text); r.font.name='Consolas'; r.font.size=Pt(9); r.font.color.rgb=RGBColor(31,65,90)
    shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),'F2F4F7'); p._p.get_or_add_pPr().append(shd)
def table(rows,widths):
    t=doc.add_table(rows=1, cols=len(rows[0])); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'; t.autofit=False
    for j,v in enumerate(rows[0]): t.rows[0].cells[j].text=v
    for row in rows[1:]:
        cells=t.add_row().cells
        for j,v in enumerate(row): cells[j].text=v
    for row in t.rows:
        for j,c in enumerate(row.cells): c.width=Inches(widths[j]); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for c in t.rows[0].cells:
        shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),'E8EEF5'); c._tc.get_or_add_tcPr().append(shd)
        for run in c.paragraphs[0].runs: run.bold=True
    return t

doc.add_paragraph().paragraph_format.space_after=Pt(70)
para('HƯỚNG DẪN TRIỂN KHAI',bold=True,align=WD_ALIGN_PARAGRAPH.CENTER)
title('CONSTRUCT-OS ERP')
para('Bản Windows và bản triển khai VPS',italic=True,align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph().paragraph_format.space_after=Pt(120)
table([['Gói','Mục đích'],['Windows','Ứng dụng desktop dành cho người dùng cuối'],['VPS','Webapp dùng chung, API và PostgreSQL']], [1.4,5.1])
para('Phiên bản tài liệu: 1.0  |  Cập nhật: 12/07/2026',align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

doc.add_heading('1. Tổng quan kiến trúc', level=1)
para('Bản Windows là ứng dụng desktop kết nối tới webapp CONSTRUCT-OS. Bản VPS là nguồn dữ liệu tập trung gồm giao diện React, API Node.js và PostgreSQL chạy bằng Docker Compose.')
table([['Thành phần','Nơi chạy','Dữ liệu'],['Windows Client','Máy người dùng','Không lưu database; kết nối VPS'],['Web/API','VPS','Xác thực, phân quyền, đồng bộ'],['PostgreSQL','Docker volume trên VPS','Dữ liệu nghiệp vụ chính'],['LocalStorage/PWA','Trình duyệt hoặc desktop','Bản đệm khi mạng yếu']], [1.45,1.45,3.6])

doc.add_heading('2. Cài bản Windows', level=1)
doc.add_heading('2.1 Yêu cầu', level=2)
bullet('Windows 10 hoặc Windows 11, 64-bit.'); bullet('Kết nối mạng tới VPS hoặc máy chạy webapp.'); bullet('Không cần cài Node.js hay trình duyệt riêng.')
doc.add_heading('2.2 Cài và chạy', level=2)
step('Giải nén CONSTRUCT-OS-ERP-Windows.zip vào thư mục cố định, ví dụ C:\\CONSTRUCT-OS.')
step('Mở desktop-config.json và thay appUrl bằng địa chỉ HTTPS của VPS.')
code('{\n  "appUrl": "https://erp.tenmiencongty.vn"\n}')
step('Chạy CONSTRUCT-OS ERP.exe. Có thể tạo shortcut ra Desktop.')
step('Đăng nhập bằng tài khoản đã được CEO cấp hoặc đăng ký bằng mã nhân viên và số điện thoại.')
doc.add_heading('2.3 Dùng với máy chủ cục bộ', level=2)
para('Nếu Docker đang chạy trên cùng máy Windows, giữ appUrl là http://localhost:8080.')

doc.add_heading('3. Triển khai VPS', level=1)
doc.add_heading('3.1 Yêu cầu máy chủ', level=2)
table([['Hạng mục','Khuyến nghị'],['Hệ điều hành','Ubuntu 22.04/24.04 LTS'],['CPU/RAM','Tối thiểu 2 vCPU, 4 GB RAM'],['Ổ đĩa','Tối thiểu 30 GB SSD'],['Phần mềm','Docker Engine và Docker Compose'],['Tên miền','Tên miền hoặc subdomain có HTTPS']], [1.7,4.8])
doc.add_heading('3.2 Cài đặt', level=2)
step('Sao chép gói CONSTRUCT-OS-ERP-VPS.zip lên VPS và giải nén.')
step('Tạo file môi trường từ mẫu.'); code('cp .env.vps.example .env\nnano .env')
step('Thay toàn bộ CHANGE_ME bằng mật khẩu ngẫu nhiên mạnh. PIN production cần 6-12 chữ số.')
step('Build và chạy hệ thống.'); code('docker compose up -d --build\ndocker compose ps\ncurl http://127.0.0.1:8080/api/health')
step('Cấu hình Nginx hoặc Caddy reverse proxy tới 127.0.0.1:8080 và bật HTTPS.')

doc.add_heading('4. Dữ liệu, đồng bộ và backup', level=1)
bullet('Dữ liệu chính nằm trong Docker volume quanlyxaydung_postgres_data.')
bullet('Không chạy docker compose down -v nếu muốn giữ dữ liệu.')
bullet('Thay đổi được lưu lên VPS sau khoảng 1,5 giây và máy khác kiểm tra phiên bản mỗi 5 giây.')
doc.add_heading('4.1 Sao lưu', level=2); code('sh scripts/backup.sh')
para('Nên chạy backup mỗi đêm bằng cron và sao chép file backup sang máy khác hoặc object storage.')
doc.add_heading('4.2 Cập nhật phiên bản', level=2)
code('sh scripts/backup.sh\ndocker compose up -d --build\ndocker compose ps')

doc.add_heading('5. Quản trị tài khoản và bảo mật', level=1)
bullet('CEO tạo, khóa, mở khóa, đặt lại PIN và xóa tài khoản nhân viên trong Vận hành Nhân sự.')
bullet('Mọi tài khoản có thể đổi PIN; PIN được băm bcrypt và không thể xem lại.')
bullet('Không đưa file .env lên GitHub hoặc gửi qua kênh không an toàn.')
bullet('Không mở cổng PostgreSQL 5432 ra Internet.')
bullet('Bắt buộc dùng HTTPS khi triển khai công khai để bảo vệ PIN và dữ liệu nhân sự.')

doc.add_heading('6. Kiểm tra sau triển khai', level=1)
for x in ['Cả app và postgres hiển thị healthy trong docker compose ps.','/api/health trả về status ok.','Đăng nhập thử CEO, Kế toán, Chỉ huy trưởng và Nhân viên.','Thêm nhân viên hoặc sửa thông tin doanh nghiệp trên máy A; máy B nhận thay đổi.','Xuất thử Excel, Word và PDF phiếu thu/chi.','Tạo backup và xác nhận file có kích thước lớn hơn 0.']: bullet(x)

doc.add_heading('7. Xử lý sự cố nhanh', level=1)
table([['Hiện tượng','Cách xử lý'],['Windows báo không kết nối được','Kiểm tra appUrl, Internet, DNS và HTTPS.'],['Webapp không mở','Chạy docker compose ps và docker compose logs -f app.'],['Database không healthy','Kiểm tra dung lượng đĩa và docker compose logs postgres.'],['Không đồng bộ','Kiểm tra trạng thái trên giao diện và /api/health.'],['Quên PIN','CEO đặt lại PIN; không thể đọc PIN đã băm.'],['Cần khôi phục','Dừng ghi dữ liệu và phục hồi từ file pg_dump gần nhất.']], [2.1,4.4])

doc.add_heading('8. Lệnh vận hành thường dùng', level=1)
code('docker compose ps\ndocker compose logs -f app\ndocker compose restart\ndocker compose down\ndocker compose up -d')
para('Cảnh báo: docker compose down -v xóa volume PostgreSQL và toàn bộ dữ liệu nếu không có backup.', bold=True)

doc.save(OUT)
print(OUT.resolve())
